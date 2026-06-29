#!/usr/bin/env python3
"""Render Brad Guillen's resume DOCX from resume_data.py (editable in Google Docs).

Usage:
    python3 scripts/generate_resume_docx.py
    python3 scripts/generate_resume_docx.py -o exports/resume.docx

This file is RENDERING ONLY — edit wording in scripts/resume_data.py.
"""
from __future__ import annotations

import argparse
import html
import os
import re
import sys
from typing import Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import resume_data as D

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x5C, 0x8A)
GRAY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x00, 0x00, 0x00)

TOKEN_RE = re.compile(
    r"(<b>|</b>|<i>|</i>|<a\s+href=\"([^\"]+)\"[^>]*>|</a>)",
    re.IGNORECASE,
)


def decode_entities(text: str) -> str:
    return html.unescape(
        text.replace("&mdash;", "—")
        .replace("&ndash;", "–")
        .replace("&nbsp;", " ")
        .replace("&bull;", "•")
    )


def tokenize_markup(text: str) -> Iterator[tuple[str, bool, bool, str | None]]:
    """Yield (text, bold, italic, link_url) segments from mini-HTML."""
    bold = False
    italic = False
    link: str | None = None
    pos = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > pos:
            chunk = decode_entities(text[pos : match.start()])
            if chunk:
                yield chunk, bold, italic, link
        token = match.group(0)
        if token == "<b>":
            bold = True
        elif token == "</b>":
            bold = False
        elif token == "<i>":
            italic = True
        elif token == "</i>":
            italic = False
        elif token.startswith("<a "):
            link = match.group(2)
        elif token == "</a>":
            link = None
        pos = match.end()
    if pos < len(text):
        chunk = decode_entities(text[pos:])
        if chunk:
            yield chunk, bold, italic, link


def set_run_font(run, *, size_pt: float, color: RGBColor, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Helvetica"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:ascii"), "Helvetica")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Helvetica")


def add_rich_paragraph(
    doc: Document,
    text: str,
    *,
    size_pt: float = 9.6,
    color: RGBColor = BLACK,
    bold: bool = False,
    italic: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_before_pt: float = 0,
    space_after_pt: float = 3,
    line_spacing: float = 1.15,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing

    if not re.search(r"<[aib/]", text):
        run = p.add_run(decode_entities(text))
        set_run_font(run, size_pt=size_pt, color=color, bold=bold, italic=italic)
        return

    for chunk, chunk_bold, chunk_italic, link in tokenize_markup(text):
        run = p.add_run(chunk)
        set_run_font(
            run,
            size_pt=size_pt,
            color=BLUE if link else color,
            bold=bold or chunk_bold,
            italic=italic or chunk_italic,
        )
        if link:
            run.underline = True
            # Hyperlink for Word/Google Docs
            part = doc.part
            r_id = part.relate_to(
                link,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            run_element = run._element
            run_element.getparent().remove(run_element)
            hyperlink.append(run_element)
            p._element.append(hyperlink)


def add_section_heading(doc: Document, title: str) -> None:
    add_rich_paragraph(
        doc,
        title,
        size_pt=10.5,
        color=NAVY,
        bold=True,
        space_before_pt=10,
        space_after_pt=2,
    )
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(4)
    p_pr = rule._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3864")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.12
    pf.left_indent = Inches(0.15)

    for chunk, chunk_bold, chunk_italic, link in tokenize_markup(decode_entities(text)):
        run = p.add_run(chunk)
        set_run_font(
            run,
            size_pt=9.5,
            color=BLUE if link else BLACK,
            bold=chunk_bold,
            italic=chunk_italic,
        )
        if link:
            run.underline = True
            part = doc.part
            r_id = part.relate_to(
                link,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            run_element = run._element
            run_element.getparent().remove(run_element)
            hyperlink.append(run_element)
            p._element.append(hyperlink)


def add_job_header(doc: Document, title_html: str, date_html: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells
    left.width = Inches(4.85)
    right.width = Inches(2.45)

    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    for chunk, chunk_bold, chunk_italic, link in tokenize_markup(decode_entities(title_html)):
        run = lp.add_run(chunk)
        set_run_font(run, size_pt=10, color=BLUE if link else BLACK, bold=True or chunk_bold, italic=chunk_italic)
        if link:
            run.underline = True

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    for chunk, chunk_bold, chunk_italic, link in tokenize_markup(decode_entities(date_html)):
        run = rp.add_run(chunk)
        set_run_font(run, size_pt=8.8, color=BLUE if link else GRAY, bold=chunk_bold, italic=chunk_italic)
        if link:
            run.underline = True


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    c = D.CONTACT
    add_rich_paragraph(doc, c["name"], size_pt=21, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=1)
    add_rich_paragraph(doc, c["title"], size_pt=11.5, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=4)

    contact_line = (
        f'{c["location"]}  |  {c["phone_display"]}  |  {c["email"]}  |  '
        f'LinkedIn: {c["linkedin"]}  |  GitHub: {c["github"]}  |  {c["website_display"]}'
    )
    add_rich_paragraph(
        doc,
        contact_line,
        size_pt=8.6,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=6,
    )

    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    add_rich_paragraph(doc, D.SUMMARY, size_pt=9.6, space_after_pt=4)

    add_section_heading(doc, "TECHNICAL PHILOSOPHY")
    add_rich_paragraph(doc, D.PHILOSOPHY, size_pt=9.6, space_after_pt=4)

    add_section_heading(doc, "SELECTED PROJECTS")
    pr = D.PROJECT
    add_job_header(
        doc,
        f'<a href="{pr["url"]}">{pr["name"]}</a> — {pr["subtitle"]}',
        f'<a href="{pr["url"]}">{pr["url_display"]}</a>',
    )
    add_rich_paragraph(doc, pr["tech"], size_pt=8.8, color=GRAY, italic=True, space_after_pt=2)
    for bullet in pr["bullets"]:
        add_bullet(doc, bullet)

    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, value in D.SKILLS:
        add_rich_paragraph(doc, f"<b>{label}:</b> {value}", size_pt=9.2, space_after_pt=1)

    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
    for role in D.EXPERIENCE:
        add_job_header(doc, role["title"], role["date"])
        if role.get("ctx"):
            add_rich_paragraph(
                doc,
                role["ctx"].format(scale=D.KW_USER_SCALE),
                size_pt=8.8,
                color=GRAY,
                italic=True,
                space_after_pt=2,
            )
        for bullet in role["bullets"]:
            add_bullet(doc, bullet)

    add_section_heading(doc, "EDUCATION & LANGUAGES")
    for line in D.EDUCATION:
        add_rich_paragraph(doc, line, size_pt=9.2, space_after_pt=1)

    return doc


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "-o",
        "--out",
        default="exports/Brad Guillen - Senior Software Engineer.docx",
    )
    args = ap.parse_args()
    os.makedirs(os.path.dirname(args.out) or ".", exist_ok=True)
    build_document().save(args.out)
    print("WROTE", args.out)


if __name__ == "__main__":
    main()
